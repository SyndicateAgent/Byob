import json
import subprocess
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document as DocxDocument
from PIL import Image

from workers.parsers import mineru_parser, registry
from workers.parsers.base import ParsedChunk, ParsedDocument
from workers.parsers.mineru_parser import MineruParserConfig, parse_mineru_document
from workers.parsers.pdf_parser import PdfParserConfig


def docx_bytes_with_image() -> tuple[bytes, bytes]:
    """Create a DOCX payload containing text and one embedded PNG image."""

    image_buffer = BytesIO()
    Image.new("RGB", (4, 3), color=(48, 96, 160)).save(image_buffer, format="PNG")
    image_content = image_buffer.getvalue()
    image_buffer.seek(0)

    document = DocxDocument()
    document.add_paragraph("Fallback DOCX paragraph")
    document.add_picture(image_buffer)
    document_buffer = BytesIO()
    document.save(document_buffer)
    return document_buffer.getvalue(), image_content


def xlsx_bytes_with_image() -> tuple[bytes, bytes]:
        """Create a minimal XLSX package with shared strings and one media asset."""

        image_buffer = BytesIO()
        Image.new("RGB", (3, 2), color=(96, 48, 160)).save(image_buffer, format="PNG")
        image_content = image_buffer.getvalue()

        workbook = """<?xml version="1.0" encoding="UTF-8"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"
                    xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
    <sheets>
        <sheet name="Summary" sheetId="1" r:id="rId1"/>
    </sheets>
</workbook>
"""
        workbook_rels = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1"
                                Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet"
                                Target="worksheets/sheet1.xml"/>
</Relationships>
"""
        shared_strings = """<?xml version="1.0" encoding="UTF-8"?>
<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
    <si><t>Quarter</t></si>
    <si><t>Revenue</t></si>
    <si><t>Q1</t></si>
</sst>
"""
        worksheet = """<?xml version="1.0" encoding="UTF-8"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
    <sheetData>
        <row r="1">
            <c r="A1" t="s"><v>0</v></c>
            <c r="B1" t="s"><v>1</v></c>
        </row>
        <row r="2">
            <c r="A2" t="s"><v>2</v></c>
            <c r="B2"><v>1250</v></c>
        </row>
    </sheetData>
</worksheet>
"""

        workbook_buffer = BytesIO()
        with ZipFile(workbook_buffer, "w") as archive:
                archive.writestr("xl/workbook.xml", workbook)
                archive.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
                archive.writestr("xl/sharedStrings.xml", shared_strings)
                archive.writestr("xl/worksheets/sheet1.xml", worksheet)
                archive.writestr("xl/media/image1.png", image_content)
        return workbook_buffer.getvalue(), image_content


@pytest.mark.parametrize("file_type", ["docx", "ppt", "pptx", "xlsx"])
def test_mineru_parser_supports_office_file_types(
    monkeypatch: pytest.MonkeyPatch,
    file_type: str,
) -> None:
    """Office documents should be parsed by MinerU with their original extension."""

    seen: dict[str, object] = {}

    def fake_run(
        command: list[str],
        *,
        check: bool,
        capture_output: bool,
        text: bool,
        timeout: int,
    ) -> subprocess.CompletedProcess[str]:
        assert check is False
        assert capture_output is True
        assert text is True
        assert timeout == 45
        input_path = Path(command[command.index("-p") + 1])
        output_dir = Path(command[command.index("-o") + 1])
        seen["input_suffix"] = input_path.suffix
        result_dir = output_dir / input_path.stem / "auto"
        result_dir.mkdir(parents=True)
        content_list = [
            {
                "type": "text",
                "text": f"{file_type.upper()} quarterly plan",
                "page_idx": 0,
            }
        ]
        (result_dir / f"{input_path.stem}_content_list.json").write_text(
            json.dumps(content_list),
            encoding="utf-8",
        )
        return subprocess.CompletedProcess(command, 0, stdout="ok", stderr="")

    monkeypatch.setattr(mineru_parser.shutil, "which", lambda command: command)
    monkeypatch.setattr(mineru_parser.subprocess, "run", fake_run)

    parsed = parse_mineru_document(
        b"office bytes",
        file_type=file_type,
        config=MineruParserConfig(timeout_seconds=45),
        source_name=f"quarterly-plan.{file_type}",
    )

    assert seen["input_suffix"] == f".{file_type}"
    assert parsed.metadata["parser"] == "mineru"
    assert parsed.metadata["mineru_file_type"] == file_type
    assert parsed.metadata["mineru_output"] == "content_list"
    assert parsed.chunks == [
        ParsedChunk(
            content=f"{file_type.upper()} quarterly plan",
            page_num=1,
            metadata={"mineru_block_type": "text"},
        )
    ]


@pytest.mark.parametrize("file_type", ["docx", "ppt", "pptx", "xlsx"])
def test_mineru_parser_extracts_office_images_and_assets(
    monkeypatch: pytest.MonkeyPatch,
    file_type: str,
) -> None:
    """Office image blocks should become image chunks with stored assets."""

    def fake_run(
        command: list[str],
        *,
        check: bool,
        capture_output: bool,
        text: bool,
        timeout: int,
    ) -> subprocess.CompletedProcess[str]:
        assert check is False
        assert capture_output is True
        assert text is True
        assert timeout == 45
        input_path = Path(command[command.index("-p") + 1])
        output_dir = Path(command[command.index("-o") + 1])
        result_dir = output_dir / input_path.stem / "auto"
        image_dir = result_dir / "images"
        image_dir.mkdir(parents=True)
        (image_dir / "chart.png").write_bytes(b"fake chart image")
        content_list = [
            {
                "type": "image",
                "caption": ["Revenue chart"],
                "img_path": "images/chart.png",
                "page_idx": 2,
                "bbox": [10, 20, 300, 420],
            }
        ]
        (result_dir / f"{input_path.stem}_content_list_v2.json").write_text(
            json.dumps(content_list),
            encoding="utf-8",
        )
        return subprocess.CompletedProcess(command, 0, stdout="ok", stderr="")

    monkeypatch.setattr(mineru_parser.shutil, "which", lambda command: command)
    monkeypatch.setattr(mineru_parser.subprocess, "run", fake_run)

    parsed = parse_mineru_document(
        b"office bytes",
        file_type=file_type,
        config=MineruParserConfig(timeout_seconds=45),
        source_name=f"quarterly-plan.{file_type}",
    )

    assert parsed.metadata["parser"] == "mineru"
    assert parsed.metadata["mineru_file_type"] == file_type
    assert parsed.metadata["mineru_block_types"] == {"image": 1}
    assert parsed.text == "Revenue chart\n![Revenue chart](images/chart.png)"
    assert len(parsed.chunks) == 1
    assert parsed.chunks[0] == ParsedChunk(
        content="Revenue chart\n![Revenue chart](images/chart.png)",
        chunk_type="image",
        page_num=3,
        bbox={"points": [10, 20, 300, 420]},
        metadata={
            "mineru_block_type": "image",
            "image_caption": "Revenue chart",
            "image_path": "images/chart.png",
        },
    )
    assert len(parsed.assets) == 1
    assert parsed.assets[0].source_path == "images/chart.png"
    assert parsed.assets[0].content == b"fake chart image"
    assert parsed.assets[0].content_type == "image/png"
    assert parsed.assets[0].metadata["aliases"] == [
        "images/chart.png",
        "./images/chart.png",
        "quarterly-plan/auto/images/chart.png",
        "./quarterly-plan/auto/images/chart.png",
    ]


@pytest.mark.parametrize("file_type", ["docx", "ppt", "pptx", "xlsx"])
def test_registry_dispatches_office_documents_to_mineru(
    monkeypatch: pytest.MonkeyPatch,
    file_type: str,
) -> None:
    """The upload parser registry should route Office documents through MinerU."""

    seen: dict[str, object] = {}

    def fake_parse_mineru_document(
        content: bytes,
        *,
        file_type: str,
        config: object,
        source_name: str | None,
    ) -> ParsedDocument:
        seen["content"] = content
        seen["file_type"] = file_type
        seen["config"] = config
        seen["source_name"] = source_name
        return ParsedDocument(
            text="office text",
            metadata={"parser": "mineru"},
            chunks=[ParsedChunk(content="office text")],
        )

    monkeypatch.setattr(registry, "parse_mineru_document", fake_parse_mineru_document)

    parsed = registry.parse_document_bytes(
        b"office bytes",
        file_type,
        pdf_config=PdfParserConfig(mineru_timeout_seconds=45),
        source_name=f"quarterly-plan.{file_type}",
    )

    assert parsed.metadata["parser"] == "mineru"
    assert parsed.chunks == [ParsedChunk(content="office text")]
    assert seen["content"] == b"office bytes"
    assert seen["file_type"] == file_type
    assert seen["source_name"] == f"quarterly-plan.{file_type}"
    assert isinstance(seen["config"], MineruParserConfig)
    assert seen["config"].timeout_seconds == 45


def test_registry_docx_fallback_extracts_embedded_images_when_mineru_is_missing(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """DOCX should prefer MinerU, but fallback parsing must still keep images."""

    monkeypatch.setattr(mineru_parser.shutil, "which", lambda command: None)
    content, image_content = docx_bytes_with_image()

    parsed = registry.parse_document_bytes(
        content,
        "docx",
        pdf_config=PdfParserConfig(mineru_command="missing-mineru"),
        source_name="fallback.docx",
    )

    assert parsed.metadata["parser"] == "python-docx"
    assert parsed.metadata["paragraph_count"] == 1
    assert parsed.metadata["image_count"] == 1
    assert parsed.metadata["mineru_fallback"] is True
    assert "MinerU command not found" in str(parsed.metadata["mineru_error"])
    assert parsed.text == "Fallback DOCX paragraph\n\n![Embedded image 1](word/media/image1.png)"
    assert [chunk.chunk_type for chunk in parsed.chunks] == ["text", "image"]
    assert parsed.chunks[0] == ParsedChunk(
        content="Fallback DOCX paragraph",
        metadata={"paragraph_index": 0},
    )
    assert parsed.chunks[1] == ParsedChunk(
        content="![Embedded image 1](word/media/image1.png)",
        chunk_type="image",
        metadata={
            "image_caption": "Embedded image 1",
            "image_path": "word/media/image1.png",
            "relationship_id": "rId9",
        },
    )
    assert len(parsed.assets) == 1
    assert parsed.assets[0].source_path == "word/media/image1.png"
    assert parsed.assets[0].content == image_content
    assert parsed.assets[0].content_type == "image/png"
    assert parsed.assets[0].metadata["aliases"] == [
        "word/media/image1.png",
        "./word/media/image1.png",
        "media/image1.png",
        "./media/image1.png",
    ]


def test_registry_xlsx_fallback_extracts_rows_and_images_when_mineru_is_missing(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """XLSX should still parse table rows and media when MinerU is unavailable."""

    monkeypatch.setattr(mineru_parser.shutil, "which", lambda command: None)
    content, image_content = xlsx_bytes_with_image()

    parsed = registry.parse_document_bytes(
        content,
        "xlsx",
        pdf_config=PdfParserConfig(mineru_command="missing-mineru"),
        source_name="fallback.xlsx",
    )

    assert parsed.metadata["parser"] == "xlsx-openxml"
    assert parsed.metadata["sheet_count"] == 1
    assert parsed.metadata["row_count"] == 2
    assert parsed.metadata["cell_count"] == 4
    assert parsed.metadata["image_count"] == 1
    assert parsed.metadata["mineru_fallback"] is True
    assert "MinerU command not found" in str(parsed.metadata["mineru_error"])
    assert "## Summary" in parsed.text
    assert "R1: Quarter | Revenue" in parsed.text
    assert "R2: Q1 | 1250" in parsed.text
    assert [chunk.chunk_type for chunk in parsed.chunks] == ["table", "image"]
    assert parsed.chunks[0] == ParsedChunk(
        content="Sheet: Summary\nR1: Quarter | Revenue\nR2: Q1 | 1250",
        chunk_type="table",
        metadata={
            "sheet_name": "Summary",
            "sheet_index": 1,
            "sheet_path": "xl/worksheets/sheet1.xml",
            "row_start": 1,
            "row_end": 2,
            "row_count": 2,
            "chunk_part": 1,
            "heading_level": 2,
        },
    )
    assert len(parsed.assets) == 1
    assert parsed.assets[0].source_path == "xl/media/image1.png"
    assert parsed.assets[0].content == image_content
    assert parsed.assets[0].content_type == "image/png"
    assert parsed.assets[0].metadata["aliases"] == [
        "xl/media/image1.png",
        "./xl/media/image1.png",
        "media/image1.png",
        "./media/image1.png",
        "../media/image1.png",
    ]