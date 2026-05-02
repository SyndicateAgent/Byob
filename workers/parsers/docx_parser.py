from io import BytesIO
from pathlib import PurePosixPath

from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RELATIONSHIP_TYPE

from workers.parsers.base import ParsedAsset, ParsedChunk, ParsedDocument


def parse_docx(content: bytes) -> ParsedDocument:
    """Extract paragraphs and embedded images from a DOCX document."""

    document = DocxDocument(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    assets, image_chunks = collect_docx_images(document)
    image_markdown = [chunk.content for chunk in image_chunks]
    return ParsedDocument(
        text="\n\n".join([*paragraphs, *image_markdown]),
        metadata={
            "parser": "python-docx",
            "paragraph_count": len(paragraphs),
            "image_count": len(assets),
        },
        assets=assets,
        chunks=[
            ParsedChunk(content=paragraph, metadata={"paragraph_index": index})
            for index, paragraph in enumerate(paragraphs)
        ]
        + image_chunks,
    )


def collect_docx_images(document: DocxDocument) -> tuple[list[ParsedAsset], list[ParsedChunk]]:
    """Collect images stored in a DOCX package as parser assets and chunks."""

    assets: list[ParsedAsset] = []
    chunks: list[ParsedChunk] = []
    for relationship in document.part.rels.values():
        if relationship.reltype != RELATIONSHIP_TYPE.IMAGE:
            continue
        image_index = len(assets) + 1
        target_ref = normalize_docx_target_ref(relationship.target_ref)
        source_path = docx_image_source_path(target_ref, image_index)
        title = f"Embedded image {image_index}"
        assets.append(
            ParsedAsset(
                source_path=source_path,
                content=relationship.target_part.blob,
                content_type=relationship.target_part.content_type,
                metadata={
                    "aliases": docx_image_aliases(source_path, target_ref),
                    "relationship_id": relationship.rId,
                },
            )
        )
        chunks.append(
            ParsedChunk(
                content=f"![{title}]({source_path})",
                chunk_type="image",
                metadata={
                    "image_caption": title,
                    "image_path": source_path,
                    "relationship_id": relationship.rId,
                },
            )
        )
    return assets, chunks


def normalize_docx_target_ref(target_ref: str) -> str:
    """Normalize a DOCX relationship target path."""

    return target_ref.replace("\\", "/").lstrip("/")


def docx_image_source_path(target_ref: str, image_index: int) -> str:
    """Return a stable parser asset path for a DOCX image relationship."""

    path = PurePosixPath(target_ref)
    if path.name:
        return f"word/{target_ref}" if not target_ref.startswith("word/") else target_ref
    return f"word/media/image{image_index}.bin"


def docx_image_aliases(source_path: str, target_ref: str) -> list[str]:
    """Return source aliases that should rewrite to the stored asset URL."""

    aliases = [source_path, f"./{source_path}"]
    if target_ref and target_ref != source_path:
        aliases.extend([target_ref, f"./{target_ref}"])
    return list(dict.fromkeys(aliases))
